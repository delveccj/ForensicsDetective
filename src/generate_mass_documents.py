#!/usr/bin/env python3
"""
Mass Document Generator for PDF Forensics Research

Generates thousands of source documents with diverse content for scaling
the PDF provenance detection research. Creates documents across multiple
academic domains with varying complexity.
"""

import os
import random
import wikipedia
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
import time

class MassDocumentGenerator:
    def __init__(self, output_dir='wikipedia_docs_expanded'):
        self.output_dir = output_dir
        self.topics = self._get_topic_categories()
        self.generated_count = 0
        self.success_count = 0
        self.fail_count = 0

        # Create output directory
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def _get_topic_categories(self):
        """Define topic categories for diverse content generation."""
        return {
            'STEM': [
                'Mathematics', 'Physics', 'Computer Science', 'Engineering',
                'Biology', 'Chemistry', 'Statistics', 'Data Science',
                'Machine Learning', 'Artificial Intelligence', 'Quantum Computing'
            ],
            'Humanities': [
                'History', 'Literature', 'Philosophy', 'Art History',
                'Linguistics', 'Anthropology', 'Archaeology', 'Ethics',
                'Political Philosophy', 'Cultural Studies', 'Religion'
            ],
            'Social Sciences': [
                'Psychology', 'Sociology', 'Economics', 'Political Science',
                'Geography', 'Demography', 'Criminology', 'Education',
                'Social Psychology', 'Urban Studies', 'International Relations'
            ],
            'Current Events': [
                'Climate Change', 'Technology Trends', 'Global Health',
                'Space Exploration', 'Renewable Energy', 'Digital Privacy',
                'Cybersecurity', 'Artificial Intelligence Ethics', 'Cryptocurrency'
            ],
            'Reference': [
                'Encyclopedia', 'Dictionary', 'Atlas', 'Almanac',
                'Technical Writing', 'Research Methods', 'Academic Publishing'
            ]
        }

    def _get_random_topic(self):
        """Get a random topic from available categories."""
        category = random.choice(list(self.topics.keys()))
        base_topic = random.choice(self.topics[category])
        
        # Try to find a real Wikipedia page
        try:
            # Search for pages in this category
            search_results = wikipedia.search(base_topic, results=10)
            if search_results:
                return random.choice(search_results).replace(' ', '_')
        except:
            pass
        
        # Fallback to synthetic
        return f"{base_topic}_{random.randint(1000, 9999)}"

    def _generate_content_complexity(self):
        """Generate document with varying complexity using Wikipedia content."""
        try:
            # Try to get real Wikipedia content
            page = wikipedia.page(self._get_random_topic().replace('_', ' '))
            content = page.content[:10000]  # Limit content length
            paragraphs = [p.strip() for p in content.split('\n') if p.strip() and not p.startswith('=')]
            if len(paragraphs) > 5:
                return paragraphs[:random.randint(5, 20)]
        except:
            pass
        
        # Fallback to synthetic content
        complexity = random.choice(['simple', 'medium', 'complex'])
        
        if complexity == 'simple':
            return self._generate_simple_content()
        elif complexity == 'medium':
            return self._generate_medium_content()
        else:
            return self._generate_complex_content()
        """Generate simple document content."""
        paragraphs = []
        num_paragraphs = random.randint(3, 8)

        for i in range(num_paragraphs):
            words = random.randint(50, 150)
            paragraph = ' '.join([f"word{j}" for j in range(words)])
            paragraphs.append(paragraph)

        return paragraphs

    def _generate_medium_content(self):
        """Generate medium complexity content with some formatting."""
        content = []

        # Add introduction
        content.append("Introduction")
        content.append("This document explores various aspects of the topic under discussion. " * random.randint(5, 15))

        # Add sections
        sections = ['Background', 'Methodology', 'Results', 'Discussion', 'Conclusion']
        for section in random.sample(sections, random.randint(2, 5)):
            content.append(f"\n{section}")
            content.append(f"This section covers {section.lower()} related information. " * random.randint(8, 20))

        return content

    def _generate_complex_content(self):
        """Generate complex content with tables, lists, etc."""
        content = []

        # Title
        content.append("Comprehensive Analysis Report")

        # Abstract
        content.append("Abstract")
        content.append("This comprehensive report provides detailed analysis of complex topics. " * random.randint(10, 25))

        # Multiple sections with subsections
        main_sections = ['Literature Review', 'Theoretical Framework', 'Empirical Analysis', 'Case Studies', 'Future Directions']

        for section in random.sample(main_sections, random.randint(3, 5)):
            content.append(f"\n{section}")

            # Subsections
            subsections = ['Overview', 'Key Findings', 'Implications', 'Limitations']
            for subsection in random.sample(subsections, random.randint(1, 3)):
                content.append(f"\n{subsection}")
                content.append(f"Detailed discussion of {subsection.lower()}. " * random.randint(15, 30))

        return content

    def _create_word_document(self, title, content, filename):
        """Create a Word document with the given content."""
        doc = Document()

        # Add title
        doc.add_heading(title.replace('_', ' '), 0)

        # Add content
        for item in content:
            if item.startswith('\n'):
                # Section header
                doc.add_heading(item[1:], level=1)
            else:
                # Regular paragraph
                doc.add_paragraph(item)

        # Save document
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def generate_single_document(self):
        """Generate a single document."""
        try:
            title = self._get_random_topic()
            content = self._generate_content_complexity()
            filename = f"{title}.docx"

            filepath = self._create_word_document(title, content, filename)

            self.success_count += 1
            print(f"[{self.generated_count:4d}] Generated: {filename}")

            return True

        except Exception as e:
            self.fail_count += 1
            print(f"[{self.generated_count:4d}] Failed: {str(e)}")
            return False

    def generate_mass_documents(self, target_count=5000):
        """Generate mass documents."""
        print(f"Starting mass document generation...")
        print(f"Target: {target_count} documents")
        print(f"Output directory: {self.output_dir}")
        print("=" * 60)

        start_time = time.time()

        while self.success_count < target_count:
            self.generated_count += 1
            self.generate_single_document()

            # Progress reporting
            if self.generated_count % 100 == 0:
                elapsed = time.time() - start_time
                rate = self.generated_count / elapsed
                eta = (target_count - self.success_count) / rate / 60  # minutes
                print(f"Progress: {self.success_count}/{target_count} ({self.success_count/target_count*100:.1f}%) - ETA: {eta:.1f} min")

        total_time = time.time() - start_time
        print("\n" + "=" * 60)
        print("Mass document generation complete!")
        print(f"Total generated: {self.generated_count}")
        print(f"Successful: {self.success_count}")
        print(f"Failed: {self.fail_count}")
        print(".2f")
        print(".2f")

def main():
    """Main execution."""
    print("Mass Document Generator for PDF Forensics")
    print("=" * 50)

    generator = MassDocumentGenerator()

    # For demonstration, generate smaller batch first
    # In production, use target_count=5000
    target = 500  # Generate more documents for scaling

    generator.generate_mass_documents(target)

    print("\nNext steps:")
    print("1. Review generated documents")
    print("2. Generate PDFs using different methods")
    print("3. Convert to binary images")
    print("4. Train classifiers")

if __name__ == "__main__":
    main()