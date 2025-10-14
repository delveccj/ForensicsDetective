#!/usr/bin/env python3
"""
Batch PDF Generation for Word and Google Docs Variations

Creates multiple PDF variations from Word documents by applying
different formatting, layouts, and content modifications.
Supports parallel processing for efficient dataset expansion.
"""

import os
import random
import time
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import threading
import subprocess
import shutil
from docx2pdf import convert
import pythoncom  # Add COM initialization
import shutil

class BatchWordPDFGenerator:
    def __init__(self, source_docs_dir='wikipedia_docs_expanded',
                 output_dir='scaled_pdfs/word_pdfs', temp_dir='temp_word_docs'):
        self.source_docs_dir = source_docs_dir
        self.output_dir = output_dir
        self.temp_dir = temp_dir

        # Create directories
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)

        # Thread safety
        self.lock = threading.Lock()

        # Get source files
        self.source_files = self._get_source_files()

    def _get_source_files(self):
        """Get list of source Word documents."""
        print(f"Looking for source docs in: {self.source_docs_dir}")
        print(f"Current working directory: {os.getcwd()}")
        print(f"Absolute path: {os.path.abspath(self.source_docs_dir)}")

        if not os.path.exists(self.source_docs_dir):
            print(f"Source directory {self.source_docs_dir} not found")
            return []

        files = [f for f in os.listdir(self.source_docs_dir) if f.endswith('.docx')]
        print(f"Found {len(files)} source Word documents")
        return files

    def _extract_content_from_docx(self, docx_path):
        """Extract paragraphs and styles from Word document."""
        try:
            doc = Document(docx_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append({
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic}
                                for run in paragraph.runs if run.text.strip()]
                    })

            return content
        except Exception as e:
            print(f"Error reading {docx_path}: {e}")
            return None

    def _apply_formatting_variation(self, doc, content, variation_type):
        """Apply different formatting variations to document."""
        # Clear existing content
        for paragraph in list(doc.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)

        # Add header if variation requires it
        if 'header' in variation_type:
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Document Analysis - {variation_type.title()}"
            header_para.style = doc.styles['Header']

        # Add footer if variation requires it
        if 'footer' in variation_type:
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated on {time.strftime('%Y-%m-%d')} - Page 1"
            footer_para.style = doc.styles['Footer']

        # Apply content with variations
        for item in content:
            para = doc.add_paragraph()

            if variation_type == 'bold':
                para.add_run(item['text']).bold = True
            elif variation_type == 'italic':
                para.add_run(item['text']).italic = True
            elif variation_type == 'mixed_formatting':
                # Apply random formatting to words
                words = item['text'].split()
                for i, word in enumerate(words):
                    run = para.add_run(word + ' ')
                    if random.random() < 0.3:  # 30% chance
                        run.bold = True
                    if random.random() < 0.2:  # 20% chance
                        run.italic = True
            elif variation_type == 'centered':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(item['text'])
            elif variation_type == 'justified':
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.add_run(item['text'])
            elif variation_type == 'indented':
                para.paragraph_format.left_indent = Inches(0.5)
                para.add_run(item['text'])
            elif variation_type == 'spaced':
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing = 1.5
                para.add_run(item['text'])
            else:  # normal
                para.add_run(item['text'])

    def _convert_docx_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF using docx2pdf with proper COM initialization."""
        try:
            # Initialize COM for this thread
            pythoncom.CoInitialize()

            # Use docx2pdf to convert the document
            convert(docx_path, pdf_path)
            return True
        except Exception as e:
            print(f"PDF conversion error with docx2pdf: {e}")
            # Fallback: Try using Microsoft Word if available (Windows)
            try:
                import win32com.client as win32
                word = win32.gencache.EnsureDispatch('Word.Application')
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                return True
            except ImportError:
                print("Microsoft Word not available for fallback conversion")
            except Exception as e:
                print(f"Word conversion fallback failed: {e}")

            # Fallback: Try using LibreOffice
            try:
                cmd = [
                    'soffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_path), docx_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    # LibreOffice saves with same name but .pdf extension
                    expected_pdf = os.path.splitext(docx_path)[0] + '.pdf'
                    if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
                        shutil.move(expected_pdf, pdf_path)
                    return True
                else:
                    print(f"LibreOffice conversion failed: {result.stderr}")
            except FileNotFoundError:
                print("Neither docx2pdf, Microsoft Word, nor LibreOffice found for PDF conversion")
                return False

        finally:
            # Uninitialize COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass

        return False

    def generate_word_variations(self, docx_filename, variations=10):
        """Generate multiple Word document variations and convert to PDF."""
        source_path = os.path.join(self.source_docs_dir, docx_filename)
        base_name = docx_filename.replace('.docx', '')

        # Extract content
        content = self._extract_content_from_docx(source_path)
        if content is None:
            return []

        generated_pdfs = []
        variation_types = [
            'normal', 'bold', 'italic', 'mixed_formatting',
            'centered', 'justified', 'indented', 'spaced',
            'header', 'footer', 'header_footer'
        ]

        for i in range(min(variations, len(variation_types))):
            try:
                variation_type = variation_types[i % len(variation_types)]

                # Create modified document
                temp_docx = os.path.join(self.temp_dir, f"{base_name}_v{i+1}_{variation_type}.docx")
                pdf_output = os.path.join(self.output_dir, f"{base_name}_word_v{i+1}_{variation_type}.pdf")

                # Load and modify document
                doc = Document(source_path)
                self._apply_formatting_variation(doc, content, variation_type)
                doc.save(temp_docx)

                # Convert to PDF
                if self._convert_docx_to_pdf(temp_docx, pdf_output):
                    generated_pdfs.append(pdf_output)
                    # Clean up temp file
                    try:
                        os.remove(temp_docx)
                    except:
                        pass
                else:
                    print(f"Failed to convert {temp_docx} to PDF")

            except Exception as e:
                print(f"Error generating Word variation {i+1} for {base_name}: {e}")

        return generated_pdfs

    def process_batch_word_generation(self, target_samples_per_type=500, max_workers=2):
        """Process batch generation of Word PDF variations."""
        print("Starting batch Word PDF generation...")
        print(f"Target samples: {target_samples_per_type}")
        print(f"Source documents: {len(self.source_files)}")
        print(f"Output directory: {self.output_dir}")
        print("=" * 60)

        if not self.source_files:
            print("No source documents found!")
            return

        # Calculate variations per document
        variations_needed = max(1, min(10, target_samples_per_type // len(self.source_files)))  # Cap at 10 variations per document
        print(f"Generating {variations_needed} variations per document")

        start_time = time.time()
        results = []

        # Process subset for demonstration
        demo_limit = min(50, len(self.source_files))  # Increased from 10 to 50
        documents_to_process = self.source_files[:demo_limit]

        print(f"Processing {len(documents_to_process)} documents (demo limit)...")

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_doc = {
                executor.submit(self.generate_word_variations, docx_file, variations_needed):
                docx_file for docx_file in documents_to_process
            }

            for future in as_completed(future_to_doc):
                docx_file = future_to_doc[future]
                try:
                    result = future.result()
                    results.append({
                        'filename': docx_file,
                        'pdfs_generated': len(result),
                        'pdf_paths': result
                    })

                    with self.lock:
                        print(f"[{len(results):2d}/{len(documents_to_process)}] Generated {len(result):2d} PDFs from {docx_file}")

                except Exception as e:
                    print(f"Exception processing {docx_file}: {e}")

        # Summary
        total_time = time.time() - start_time
        total_pdfs = sum(r['pdfs_generated'] for r in results)

        print("\n" + "=" * 60)
        print("BATCH WORD PDF GENERATION COMPLETE")
        print("=" * 60)
        print(f"Documents processed: {len(results)}")
        print(f"Total PDFs generated: {total_pdfs}")
        print(".2f")
        print(".2f")

        return results

class BatchGoogleDocsPDFGenerator:
    """Placeholder for Google Docs batch generation."""
    def __init__(self, output_dir='scaled_pdfs/google_docs_pdfs'):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_batch_google_docs_pdfs(self, target_samples=500):
        """Generate Google Docs PDFs (placeholder implementation)."""
        print("Google Docs batch generation not yet implemented")
        print("This would require Google Docs API integration")
        print(f"Target samples: {target_samples}")
        print("Output directory:", self.output_dir)

        # For now, return empty results
        return []

def main():
    """Main execution for batch PDF generation."""
    print("Batch PDF Generation for Dataset Expansion")
    print("=" * 50)

    # Check dependencies
    try:
        import docx
        print("[OK] python-docx available")
    except ImportError:
        print("[ERROR] python-docx not installed. Run: pip install python-docx")
        return

    print()

    # Generate Word PDFs
    print("Generating Word document variations...")
    word_generator = BatchWordPDFGenerator()
    word_results = word_generator.generate_batch_word_pdfs(target_samples=10000)

    # Generate Google Docs PDFs (placeholder)
    print("\nGenerating Google Docs variations...")
    google_generator = BatchGoogleDocsPDFGenerator()
    google_results = google_generator.generate_batch_google_docs_pdfs(target_samples=10000)

    print("\n" + "=" * 60)
    print("BATCH GENERATION SUMMARY")
    print("=" * 60)
    word_count = sum(r['pdfs_generated'] for r in word_results) if word_results else 0
    print(f"Word PDFs generated: {word_count}")
    print(f"Google Docs PDFs generated: {len(google_results)}")
    print(f"Total batch PDFs: {word_count + len(google_results)}")

    print("\nNext steps:")
    print("1. Implement Google Docs API integration")
    print("2. Scale to full 5000+ samples per type")
    print("3. Convert all PDFs to binary images")
    print("4. Validate dataset quality and classifier performance")

if __name__ == "__main__":
    main()