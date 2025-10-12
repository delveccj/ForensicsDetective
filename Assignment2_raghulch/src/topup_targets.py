import csv, random, textwrap
from pathlib import Path
import numpy as np
from faker import Faker
from tqdm import trange
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import wikipediaapi


USER_AGENT = "ForensicsDetective/1.0 (GitHub: raghuldav; contact: raghuldav@gmail.com)"

ROOT = Path(__file__).resolve().parents[1]
SDOC = ROOT/"data"/"source_documents"
DOCX = SDOC/"docx"; HTML = SDOC/"html"; ASSETS = SDOC/"assets"; META = SDOC/"_meta.csv"
for p in [DOCX, HTML, ASSETS]: p.mkdir(parents=True, exist_ok=True)

fake = Faker()
wiki = wikipediaapi.Wikipedia(
    language="en",
    user_agent=USER_AGENT,
)

def wiki_text(title):
    p = wiki.page(title)
    if not p.exists(): return title + " — " + " ".join(fake.sentences(nb=30))
    txt = p.summary or ""
    if len(txt) < 1400: txt += "\n\n" + "\n".join([s.text for s in p.sections[:3]])
    if len(txt) < 1000: txt += "\n\n" + " ".join(fake.sentences(nb=60))
    return txt

def make_chart(path):
    x = np.arange(1, 11); y = np.cumsum(np.random.randn(10)) + 10
    import matplotlib.pyplot as plt
    plt.figure(); plt.plot(x,y); plt.title("Synthetic Chart")
    plt.savefig(path, bbox_inches="tight"); plt.close()

def add_table(doc, rows=6, cols=4):
    t = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            t.cell(r,c).text = f"{fake.word()} {random.randint(1,999)}"

def chunk_text(doc, text, pages):
    words = textwrap.wrap(text, width=90)
    parts = np.array_split(words, max(1, pages*8))
    for part in parts: doc.add_paragraph(" ".join(part))

def build(idx, title, pages):
    d = Document()
    d.add_heading(title, 0)
    for lvl in [1,2]:
        d.add_heading(fake.sentence(nb_words=4), lvl)
        chunk_text(d, wiki_text(title), max(1, pages-1))
        d.add_paragraph()
        add_table(d, rows=6, cols=4)
        chart = ASSETS/f"chart_{idx}.png"; make_chart(chart)
        d.add_picture(str(chart), width=Inches(5))
        d.add_paragraph("Key points:")
        for _ in range(3): d.add_paragraph(f"- {fake.sentence()}")
        d.add_paragraph(f"[1] {fake.name()}, {fake.company()} ({random.randint(2001,2024)})")
        d.add_paragraph(f"Footnote: {fake.sentence()}")
    name = f"{idx:06d}_{title.replace(' ','_')}"
    d.save(DOCX/f"{name}.docx")
    (HTML/f"{name}.html").write_text(f"<h1>{title}</h1><p>{wiki_text(title)}</p>", encoding="utf-8")

def main():
    rows = []
    if META.exists():
        with open(META) as f: rows = list(csv.DictReader(f))
    next_id = (max([int(x["id"]) for x in rows]) + 1) if rows else 1

    import pandas as pd
    df = pd.DataFrame(rows) if rows else pd.DataFrame(columns=["pages_target"])
    df["pages_target"] = pd.to_numeric(df["pages_target"], errors="coerce").fillna(0).astype(int)
    b15_30 = (df["pages_target"].between(15,30)).sum()
    b30p   = (df["pages_target"]>=30).sum()

    need_15_30 = max(0, 200 - b15_30)
    need_30p   = max(0,  50 - b30p)

    topics = ["Algebra","Quantum mechanics","Machine learning","World War II","Renaissance",
              "Photosynthesis","Blockchain","Psychology","Sociology","Economics","Virtual reality",
              "United States","United Nations","Urbanization","Veganism","Volcano","Weather","Statistics"]

    added = []
    for _ in trange(need_15_30, desc="Top-up 15–30"):
        title = random.choice(topics); pages = random.randint(15,30)
        build(next_id, title, pages)
        added.append(dict(id=next_id, title=title, pages_target=pages, has_images=True, has_tables=True))
        next_id += 1
    for _ in trange(need_30p, desc="Top-up 30+"):
        title = random.choice(topics); pages = random.randint(30,45)
        build(next_id, title, pages)
        added.append(dict(id=next_id, title=title, pages_target=pages, has_images=True, has_tables=True))
        next_id += 1

    mode = "a" if META.exists() else "w"
    with open(META, mode, newline="") as f:
        w = csv.DictWriter(f, fieldnames=["id","title","pages_target","has_images","has_tables"])
        if mode=="w": w.writeheader()
        w.writerows(added)
    print(f"Top-up added {len(added)} docs.")
if __name__ == "__main__":
    main()
