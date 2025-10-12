import random, textwrap
from pathlib import Path
import numpy as np
import wikipediaapi
from faker import Faker
from tqdm import trange
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT/"data/source_documents/docx"
OUT_HTML = ROOT/"data/source_documents/html"
ASSETS   = ROOT/"data/source_documents/assets"
META     = ROOT/"data/source_documents/_meta.csv"
for p in [OUT_DOCX, OUT_HTML, ASSETS]:
    p.mkdir(parents=True, exist_ok=True)

fake = Faker()
wiki = wikipediaapi.Wikipedia(
    user_agent="ForensicsDetective-Assignment3 (https://github.com/raghuldav/ForensicsDetective; raghuldav@users.noreply.github.com) v1",
    language="en",
)

# seed topics (use pdf_names.txt if available; else fallback list)
topics = []
pdf_names = (ROOT.parent/"ForensicsDetective"/"pdf_names.txt")
if pdf_names.exists():
    topics = [t.strip().replace(".pdf","").replace("_"," ") for t in pdf_names.read_text().splitlines() if t.strip()]
if not topics:
    topics = ["Algebra","Quantum mechanics","Machine learning","World War II","Renaissance",
              "Photosynthesis","Blockchain","Psychology","Sociology","Economics","Virtual reality",
              "United States","United Nations","Urbanization","Veganism","Volcano","Weather","Statistics"]

while len(topics) < 12000:
    topics += random.sample(topics, k=min(1000, len(topics)))

def wiki_text(title):
    p = wiki.page(title)
    if not p.exists():
        return title + " — " + " ".join(fake.sentences(nb=30))
    txt = p.summary or ""
    if len(txt) < 1400:
        txt += "\n\n" + "\n".join([s.text for s in p.sections[:3]])
    if len(txt) < 1000:
        txt += "\n\n" + " ".join(fake.sentences(nb=60))
    return txt

def target_pages(mean=3, sd=1.5, minp=1, maxp=40):
    return max(minp, min(int(np.random.normal(mean, sd)), maxp))

def make_chart(path):
    x = np.arange(1, 11); y = np.cumsum(np.random.randn(10)) + 10
    plt.figure(); plt.plot(x, y); plt.title("Synthetic Chart")
    plt.savefig(path, bbox_inches="tight"); plt.close()

def add_table(doc, rows=6, cols=4):
    t = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            t.cell(r,c).text = f"{fake.word()} {random.randint(1,999)}"

def chunk_text(doc, text, pages):
    words = textwrap.wrap(text, width=90)
    parts = np.array_split(words, max(1, pages*8))
    for part in parts:
        doc.add_paragraph(" ".join(part))

def build_one(idx, title, pages):
    has_img = random.random() < 0.35
    has_tbl = random.random() < 0.30

    d = Document()
    d.add_heading(title, 0)
    for lvl in [1,2]:
        d.add_heading(fake.sentence(nb_words=4), lvl)
        chunk_text(d, wiki_text(title), max(1, pages-1))
        d.add_paragraph()
        if has_tbl and random.random() < 0.7:
            add_table(d, rows=6, cols=4)
        if has_img:
            chart = ASSETS/f"chart_{idx}.png"; make_chart(chart)
            d.add_picture(str(chart), width=Inches(5))
        d.add_paragraph("Key points:")
        for _ in range(3): d.add_paragraph(f"- {fake.sentence()}")
        d.add_paragraph(f"[1] {fake.name()}, {fake.company()} ({random.randint(2001,2024)})")
        d.add_paragraph(f"Footnote: {fake.sentence()}")

    name = f"{idx:06d}_{title.replace(' ','_')}"
    d.save(OUT_DOCX/f"{name}.docx")

    html = [f"<h1>{title}</h1>"]
    for _ in range(3):
        html += [f"<h2>{fake.sentence()}</h2>", f"<p>{wiki_text(title)}</p>", "<ul>"]
        for __ in range(3): html.append(f"<li>{fake.sentence()}</li>")
        html.append("</ul>")
    (OUT_HTML/f"{name}.html").write_text("\n".join(html), encoding="utf-8")

    return has_img, has_tbl, pages

def main(N=6000):
    import csv
    META.parent.mkdir(parents=True, exist_ok=True)
    with open(META, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["id","title","pages_target","has_images","has_tables"])
        w.writeheader()
        for i in trange(N, desc="Generating"):
            t = topics[i % len(topics)]
            pg = target_pages()
            has_img, has_tbl, _ = build_one(i+1, t, pg)
            w.writerow(dict(id=i+1, title=t, pages_target=pg,
                            has_images=has_img, has_tables=has_tbl))
if __name__ == "__main__":
    main(6000)
