import csv, random
from pathlib import Path
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
SDOC = ROOT/"data/source_documents"
DOCX = SDOC/"docx"; HTML = SDOC/"html"; ASSETS = SDOC/"assets"; META = SDOC/"_meta.csv"
for p in [DOCX, HTML, ASSETS]: p.mkdir(parents=True, exist_ok=True)

def make_chart(path):
    x = np.arange(1, 11); y = np.cumsum(np.random.randn(10)) + 10
    plt.figure(); plt.plot(x,y); plt.title("Chart"); plt.savefig(path, bbox_inches="tight"); plt.close()

def add_table(doc, rows=6, cols=4):
    t = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            t.cell(r,c).text = f"cell {r},{c}"

def add_body(doc, pages):
    # enough paragraphs to approx. fill N pages downstream
    words = "lorem ipsum dolor sit amet consectetur adipiscing elit "*2000
    w = words.split()
    parts = np.array_split(w, max(1, pages*8))
    for p in parts:
        doc.add_paragraph(" ".join(p))

def build(idx, title, pages):
    d = Document()
    d.add_heading(title, 0)
    d.add_heading("Section", 1)
    add_body(d, pages)
    add_table(d)
    img = ASSETS/f"tail_{idx}.png"; make_chart(img)
    d.add_picture(str(img), width=Inches(5))
    d.add_paragraph("Key points:"); d.add_paragraph("- item 1"); d.add_paragraph("- item 2")
    d.add_paragraph("[1] Example, Co. (2024)"); d.add_paragraph("Footnote: sample.")
    name = f"{idx:06d}_{title.replace(' ','_')}"
    d.save(DOCX/f"{name}.docx")
    (HTML/f"{name}.html").write_text(f"<h1>{title}</h1><p>tail filler</p>", encoding="utf-8")

def main():
    df = pd.read_csv(META) if META.exists() else pd.DataFrame(columns=["id","title","pages_target","has_images","has_tables"])
    if df.empty:
        next_id = 1
        c15_30 = 0; c30p = 0
    else:
        next_id = int(df["id"].max()) + 1
        pages = pd.to_numeric(df["pages_target"], errors="coerce").fillna(1).astype(int)
        c15_30 = ((pages>=15)&(pages<=30)).sum()
        c30p   = (pages>=30).sum()

    need_15_30 = max(0, 200 - c15_30)
    need_30p   = max(0,  50 - c30p)

    rows = []
    for _ in range(need_15_30):
        pages = random.randint(15,30)
        build(next_id, "Tail_15_30", pages)
        rows.append(dict(id=next_id, title="Tail_15_30", pages_target=pages, has_images=1, has_tables=1))
        next_id += 1
    for _ in range(need_30p):
        pages = random.randint(30,45)
        build(next_id, "Tail_30_plus", pages)
        rows.append(dict(id=next_id, title="Tail_30_plus", pages_target=pages, has_images=1, has_tables=1))
        next_id += 1

    mode = "a" if META.exists() else "w"
    fields = ["id","title","pages_target","has_images","has_tables"]
    with open(META, mode, newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        if mode=="w": w.writeheader()
        w.writerows(rows)
    print(f"Added tail docs — 15–30: {need_15_30}, 30+: {need_30p}")
if __name__ == "__main__":
    main()
